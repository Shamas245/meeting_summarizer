import streamlit as st
import os
import logging
import tempfile
import time
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv
import whisper
import google.generativeai as genai
from moviepy.editor import VideoFileClip
from docx import Document
from io import BytesIO
import traceback
import librosa
import numpy as np

try:
    from config import (
        SUMMARIZATION_PROMPT, ACTION_ITEMS_PROMPT, MEETING_TYPE_PROMPTS,
        MAX_FILE_SIZE_MB, ERROR_MESSAGES, SUCCESS_MESSAGES, UI_CONFIG,
        WHISPER_MODEL, GEMINI_MODEL, AUDIO_CONFIG
    )
    from slack import send_to_slack, validate_webhook_url, test_slack_connection
except ImportError as e:
    st.error(f"Missing required files: {e}. Please ensure config.py and slack.py exist.")
    st.stop()

# Setup logging
logging.basicConfig(
    level=logging.INFO, 
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler('app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class MeetingSummarizer:
    def __init__(self):
        self.load_environment()
        self.initialize_models()

    def load_environment(self):
        """Load and validate environment variables"""
    # Use Streamlit secrets instead of dotenv
        try:
            self.gemini_api_key = st.secrets["GEMINI_API_KEY"]
        except KeyError:
            st.error("❌ GEMINI_API_KEY not found in Streamlit secrets.")
            st.stop()
            
    def initialize_models(self):
        """Initialize ML models with error handling"""
        # Initialize Whisper model
        try:
            with st.spinner("Loading speech recognition model..."):
                self.stt_model = whisper.load_model(WHISPER_MODEL)
            logger.info(f"Whisper model '{WHISPER_MODEL}' loaded successfully")
        except Exception as e:
            st.error(f"❌ Failed to load Whisper model: {str(e)}")
            logger.error(f"Whisper loading error: {e}")
            st.stop()

        # Initialize Gemini
        try:
            genai.configure(api_key=self.gemini_api_key)
            self.gemini_model = genai.GenerativeModel(GEMINI_MODEL)
            logger.info(f"Gemini model '{GEMINI_MODEL}' initialized successfully")
        except Exception as e:
            st.error(f"❌ Failed to initialize Gemini API: {str(e)}")
            logger.error(f"Gemini init error: {e}")
            st.stop()
    
    def validate_file_size(self, file, max_size_mb=MAX_FILE_SIZE_MB):
        """Validate file size"""
        max_size_bytes = max_size_mb * 1024 * 1024
        if file.size > max_size_bytes:
            st.error(ERROR_MESSAGES["file_too_large"].format(max_size=max_size_mb))
            return False
        return True
    
    def extract_audio_from_video(self, video_file):
        """Extract audio from video file"""
        try:
            # Create temporary files using tempfile module for better cleanup
            with tempfile.NamedTemporaryFile(delete=False, suffix='.mp4') as tmp_video:
                tmp_video.write(video_file.read())
                tmp_video_path = tmp_video.name
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as tmp_audio:
                tmp_audio_path = tmp_audio.name
            
            logger.info(f"Processing video: {tmp_video_path}")
            
            # Extract audio using moviepy
            video = VideoFileClip(tmp_video_path)
            
            if video.audio is None:
                raise Exception("No audio track found in video file")
            
            # Write audio file
            video.audio.write_audiofile(tmp_audio_path, verbose=False, logger=None)
            video.audio.close()
            video.close()
            
            # Verify audio file exists and has content
            if not os.path.exists(tmp_audio_path) or os.path.getsize(tmp_audio_path) == 0:
                raise Exception("Failed to extract audio from video")
            
            logger.info(f"Audio extracted successfully: {tmp_audio_path}")
            return tmp_video_path, tmp_audio_path
            
        except Exception as e:
            logger.error(f"Audio extraction error: {e}")
            # Cleanup on error
            for path in [tmp_video_path, tmp_audio_path]:
                if 'path' in locals() and os.path.exists(path):
                    os.unlink(path)
            raise
    
    def transcribe_audio(self, audio_path):
        """Transcribe audio using Whisper with librosa for audio loading"""
        try:
            logger.info(f"Transcribing audio: {audio_path}")
            
            # Verify file exists and is readable
            if not os.path.exists(audio_path):
                raise Exception("Audio file not found")
            
            if os.path.getsize(audio_path) == 0:
                raise Exception("Audio file is empty")
            
            # Add a small delay to ensure file is fully written
            time.sleep(1)
            
            # Load audio with librosa instead of Whisper's built-in loader
            audio_array, sr = librosa.load(audio_path, sr=16000, mono=True)
            
            # Pass the audio array directly to Whisper
            result = self.stt_model.transcribe(audio_array)
            transcript = result["text"].strip()
            
            if not transcript:
                raise Exception("No speech detected in audio")
            
            logger.info("Transcription completed successfully")
            return transcript
            
        except Exception as e:
            logger.error(f"Transcription error: {e}")
            raise
    
    def generate_summary_and_actions(self, transcript, meeting_type="general"):
        """Generate summary and action items using Gemini"""
        try:
            if not transcript.strip():
                raise Exception("Transcript is empty")
            
            # Select appropriate prompts based on meeting type
            if meeting_type in MEETING_TYPE_PROMPTS:
                summary_prompt = MEETING_TYPE_PROMPTS[meeting_type]["summary"].format(transcript=transcript)
                action_prompt = MEETING_TYPE_PROMPTS[meeting_type]["actions"].format(transcript=transcript)
            else:
                summary_prompt = SUMMARIZATION_PROMPT.format(transcript=transcript)
                action_prompt = ACTION_ITEMS_PROMPT.format(transcript=transcript)
            
            # Generate summary
            logger.info(f"Generating summary for {meeting_type} meeting...")
            summary_response = self.gemini_model.generate_content(summary_prompt)
            summary = summary_response.text.strip()
            
            # Generate action items
            logger.info("Generating action items...")
            action_response = self.gemini_model.generate_content(action_prompt)
            action_text = action_response.text.strip()
            
            # Parse action items (improved parsing)
            action_items = []
            for line in action_text.split('\n'):
                line = line.strip()
                if line and (line.startswith('-') or line.startswith('•') or line.startswith('*')):
                    action_items.append(line)
                elif line and len(action_items) == 0 and not line.lower().startswith('no'):  # First item might not have bullet
                    action_items.append(f"- {line}")
            
            logger.info(f"Generated summary and {len(action_items)} action items")
            return summary, action_items
            
        except Exception as e:
            logger.error(f"Summary generation error: {e}")
            raise
    
    def create_docx(self, transcript, summary, action_items):
        """Create DOCX document"""
        try:
            doc = Document()
            
            # Title and metadata
            doc.add_heading("Meeting Summary Report", 0)
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("─" * 50)
            
            # Summary section
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(summary)
            
            # Action items section
            if action_items:
                doc.add_heading("Action Items", level=1)
                for item in action_items:
                    doc.add_paragraph(item, style="List Bullet")
            
            # Full transcript section
            doc.add_heading("Full Transcript", level=1)
            doc.add_paragraph(transcript)
            
            # Save to BytesIO
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            logger.info("DOCX document created successfully")
            return docx_buffer
            
        except Exception as e:
            logger.error(f"DOCX creation error: {e}")
            raise

def init_session_state():
    """Initialize session state variables"""
    defaults = {
        "transcript": "",
        "summary": "",
        "action_items": [],
        "docx_file": None,
        "processing_complete": False
    }
    
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

def display_header():
    """Display application header and info"""
    st.title("📝 Meeting Summarizer")
    st.markdown(f"""
    Upload your meeting content to generate professional summaries and action items.
    
    **Supported formats:**
    - 🎥 Video: {', '.join(AUDIO_CONFIG['supported_video_formats']).upper()} (max {MAX_FILE_SIZE_MB}MB)  
    - 🎵 Audio: {', '.join(AUDIO_CONFIG['supported_audio_formats']).upper()} (max {MAX_FILE_SIZE_MB}MB)
    - 📄 Text: TXT files
    """)

def display_meeting_type_selector():
    """Display meeting type selector"""
    st.subheader("📋 Meeting Type")
    meeting_type = st.selectbox(
        "Select meeting type for optimized analysis:",
        options=["general", "standup", "planning", "retrospective"],
        format_func=lambda x: {
            "general": "🎯 General Meeting",
            "standup": "⚡ Daily Standup/Scrum", 
            "planning": "📊 Planning Session",
            "retrospective": "🔄 Retrospective"
        }[x],
        help="Choose the type that best matches your meeting for more targeted analysis"
    )
    return meeting_type

def handle_file_uploads(summarizer):
    """Handle file uploads and return selected file info"""
    # Create tabs for different input methods
    tab1, tab2, tab3 = st.tabs(["🎥 Video Upload", "🎵 Audio Upload", "📄 Text Upload"])
    
    uploaded_file = None
    file_type = None
    
    # Video upload tab
    with tab1:
        video_file = st.file_uploader(
            "Upload Video File", 
            type=AUDIO_CONFIG["supported_video_formats"],
            help=f"Upload your meeting recording ({', '.join(AUDIO_CONFIG['supported_video_formats']).upper()} format, max {MAX_FILE_SIZE_MB}MB)"
        )
        if video_file and summarizer.validate_file_size(video_file):
            uploaded_file = video_file
            file_type = "video"
            st.success(SUCCESS_MESSAGES["file_uploaded"])
    
    # Audio upload tab
    with tab2:
        audio_file = st.file_uploader(
            "Upload Audio File", 
            type=AUDIO_CONFIG["supported_audio_formats"],
            help=f"Upload your meeting audio ({', '.join(AUDIO_CONFIG['supported_audio_formats']).upper()} format, max {MAX_FILE_SIZE_MB}MB)"
        )
        if audio_file and summarizer.validate_file_size(audio_file):
            uploaded_file = audio_file
            file_type = "audio"
            st.success(SUCCESS_MESSAGES["file_uploaded"])
    
    # Text upload tab
    with tab3:
        transcript_file = st.file_uploader(
            "Upload Transcript", 
            type=["txt"],
            help="Upload a text file containing your meeting transcript"
        )
        if transcript_file:
            try:
                content = transcript_file.read().decode("utf-8").strip()
                if content:
                    st.session_state.transcript = content
                    uploaded_file = transcript_file
                    file_type = "text"
                    st.success(SUCCESS_MESSAGES["file_uploaded"])
                else:
                    st.error(ERROR_MESSAGES["empty_transcript"])
            except Exception as e:
                st.error(f"❌ Failed to read transcript: {str(e)}")
    
    return uploaded_file, file_type

def process_meeting(summarizer, uploaded_file, file_type, meeting_type):
    """Process the uploaded meeting file"""
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    try:
        # Step 1: Handle file processing
        if file_type in ["video", "audio"] and not st.session_state.transcript:
            temp_files = []
            
            try:
                if file_type == "video":
                    status_text.text("🎬 Extracting audio from video...")
                    progress_bar.progress(10)
                    
                    tmp_video_path, tmp_audio_path = summarizer.extract_audio_from_video(uploaded_file)
                    temp_files.extend([tmp_video_path, tmp_audio_path])
                    audio_path = tmp_audio_path
                    
                else:  # audio file
                    status_text.text("🎵 Processing audio file...")
                    progress_bar.progress(10)
                    
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as tmp_audio:
                        tmp_audio.write(uploaded_file.read())
                        audio_path = tmp_audio.name
                        temp_files.append(audio_path)
                
                # Step 2: Transcription
                status_text.text("🎤 Converting speech to text...")
                progress_bar.progress(40)
                
                st.session_state.transcript = summarizer.transcribe_audio(audio_path)
                
            finally:
                # Cleanup temporary files
                for temp_file in temp_files:
                    if os.path.exists(temp_file):
                        try:
                            os.unlink(temp_file)
                        except Exception as e:
                            logger.warning(f"Failed to cleanup {temp_file}: {e}")
        
        # Step 3: Generate summary and action items
        if st.session_state.transcript:
            status_text.text("🤖 Generating summary and action items...")
            progress_bar.progress(70)
            
            summary, action_items = summarizer.generate_summary_and_actions(
                st.session_state.transcript, 
                meeting_type
            )
            st.session_state.summary = summary
            st.session_state.action_items = action_items
            
            # Step 4: Create DOCX
            status_text.text("📄 Creating document...")
            progress_bar.progress(90)
            
            st.session_state.docx_file = summarizer.create_docx(
                st.session_state.transcript,
                st.session_state.summary,
                st.session_state.action_items
            )
            
            progress_bar.progress(100)
            status_text.text(SUCCESS_MESSAGES["processing_complete"])
            st.session_state.processing_complete = True
            
            time.sleep(1)
            progress_bar.empty()
            status_text.empty()
            
            st.success("🎉 Meeting processed successfully!")
            
    except Exception as e:
        progress_bar.empty()
        status_text.empty()
        st.error(f"❌ Processing failed: {str(e)}")
        logger.error(f"Processing error: {traceback.format_exc()}")

def display_results():
    """Display processing results"""
    if not st.session_state.processing_complete:
        return
        
    st.markdown("---")
    
    # Summary section
    if st.session_state.summary:
        st.subheader("📋 Meeting Summary")
        st.markdown(st.session_state.summary)
    
    # Action items section
    if st.session_state.action_items:
        st.subheader("✅ Action Items")
        for i, item in enumerate(st.session_state.action_items, 1):
            st.markdown(f"{i}. {item}")
    
    # Transcript section (collapsible)
    if st.session_state.transcript:
        with st.expander("📝 View Full Transcript"):
            # Limit display length for performance
            display_transcript = st.session_state.transcript
            if len(display_transcript) > UI_CONFIG.get("max_transcript_display", 1000):
                display_transcript = display_transcript[:UI_CONFIG.get("max_transcript_display", 1000)] + "..."
            
            st.text_area("", display_transcript, height=300, disabled=True)

def display_download_and_share_options(summarizer):
    """Display download and sharing options"""
    if not st.session_state.processing_complete:
        return
        
    st.markdown("---")
    col1, col2 = st.columns(2)
    
    # Download DOCX
    with col1:
        if st.session_state.docx_file:
            st.subheader("💾 Download")
            filename = f"meeting_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            st.download_button(
                label="📄 Download Summary (DOCX)",
                data=st.session_state.docx_file,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    
    # Send to Slack
    with col2:
        if st.session_state.action_items:
            st.subheader("📤 Share")
            # Check if user provided webhook
            if 'slack_webhook' not in st.session_state:
                st.info("💡 Enter a Slack webhook URL above to enable Slack sharing")
                st.button("📱 Send to Slack", disabled=True, use_container_width=True)
                return
            
            if st.button("📱 Send to Slack", use_container_width=True):
                try:
                    with st.spinner("Sending to Slack..."):
                        success = send_to_slack(
                            summary=st.session_state.summary,
                            action_items="\n".join(st.session_state.action_items),
                            webhook_url=st.session_state['slack_webhook']
                        )
                    
                    if success:
                        st.success(SUCCESS_MESSAGES["slack_sent"])
                    else:
                        st.error(ERROR_MESSAGES["slack_error"])
                        
                except Exception as e:
                    st.error(f"❌ Slack error: {str(e)}")
                    logger.error(f"Slack send error: {e}")

def apply_custom_css():
    """Apply custom CSS styling"""
    st.markdown("""
    <style>
        .stButton > button {
            border-radius: 8px;
            border: none;
            padding: 0.5rem 1rem;
            font-weight: 600;
            transition: all 0.3s ease;
        }
        
        .stButton > button[kind="primary"] {
            background: linear-gradient(45deg, #FF6B6B, #4ECDC4);
            color: white;
        }
        
        .stButton > button:hover {
            transform: translateY(-2px);
            box-shadow: 0 4px 12px rgba(0,0,0,0.15);
        }
        
        .stProgress > div > div {
            background: linear-gradient(45deg, #FF6B6B, #4ECDC4);
        }
        
        .stTabs [data-baseweb="tab-list"] {
            gap: 8px;
        }
        
        .stTabs [data-baseweb="tab"] {
            border-radius: 8px;
            padding: 8px 16px;
        }
        
        div[data-testid="metric-container"] {
            background-color: #f0f2f6;
            border: 1px solid #d0d2d6;
            padding: 1rem;
            border-radius: 8px;
        }
        
        .upload-info {
            background: linear-gradient(135deg, rgba(78, 205, 196, 0.1), rgba(255, 107, 107, 0.1));
            padding: 1rem;
            border-radius: 8px;
            border-left: 4px solid #4ECDC4;
            margin: 1rem 0;
        }
    </style>
    """, unsafe_allow_html=True)

def main():
    # Page configuration
    st.set_page_config(
        page_title="Meeting Summarizer",
        page_icon="📝",
        layout="wide",
        initial_sidebar_state="collapsed"
    )
    
    # Apply custom styling
    apply_custom_css()
    
    # Initialize session state
    init_session_state()
    
    # Initialize the summarizer
    try:
        summarizer = MeetingSummarizer()
    except Exception as e:
        st.error(f"Failed to initialize application: {e}")
        return
    
    # Display main UI components
    display_header()
    meeting_type = display_meeting_type_selector()
    uploaded_file, file_type = handle_file_uploads(summarizer)
    
    # Add Slack webhook input
    slack_url = st.text_input("🔗 Slack Webhook URL", 
                            type="password", 
                            help="Paste your Slack webhook URL to send results to Slack")

    # Add instructions expander right below the input
    with st.expander("📘 How to get your Slack Webhook URL?"):
        st.markdown("""
        1. Go to [Slack API: Incoming Webhooks](https://api.slack.com/messaging/webhooks).
        2. Click **Create Your Own App**.
        3. Give your app a name (e.g., `MeetingBot`) and choose your workspace.
        4. Under **Add features and functionality**, click **Incoming Webhooks**.
        5. Activate **Incoming Webhooks**.
        6. Scroll down and click **Add New Webhook to Workspace**.
        7. Select a channel you want to post to.
        8. Click **Allow**.
        9. Copy the **Webhook URL** that looks like:
        ```
        https://hooks.slack.com/services/T00000000/B00000000/XXXXXXXXXXXXXXXXXXXXXXXX
        ```
        10. Paste it in the Slack Webhook field above.
        """)

    if slack_url:
        st.session_state['slack_webhook'] = slack_url
    # Process button
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        process_button = st.button(
            "🚀 Process Meeting", 
            disabled=not uploaded_file,
            use_container_width=True,
            type="primary"
        )
    
    # Processing logic
    if process_button and uploaded_file:
        process_meeting(summarizer, uploaded_file, file_type, meeting_type)
    
    # Display results and options
    display_results()
    display_download_and_share_options(summarizer)

if __name__ == "__main__":
    main()